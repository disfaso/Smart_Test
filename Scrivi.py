from docx import Document
from Verifica import Verifica
import os
from docx.shared import RGBColor

def Cartella_controllo(cartella_path):
    """
    Crea una cartella se non esiste già.

    Args:
        cartella_path (str): Il percorso della cartella da controllare/creare.

    Returns:
        None
    """
    if not os.path.exists(cartella_path):
        os.makedirs(cartella_path)

    else:
        pass

def Colore_rosso(paragraph):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 0, 0)


def Scrivi_verifica_s(verifica: Verifica, classificazione: str):
    """
    Scrive una verifica per studenti in un file .docx.

    Args:
        verifica (Verifica): Oggetto Verifica contenente gli esercizi della verifica.
        classificazione (str): La classificazione della verifica.

    Returns:
        None
    """
    tematica = verifica.esercizi[0].tematica
    sottotematica = verifica.esercizi[0].argomento.sottotematica
    materia = verifica.esercizi[0].materia
    i = 1
    file_path = f"./Smart_Test/Verifiche_{materia}/"
   

    doc = Document()

    doc.add_heading(f"Verifica di {materia}; unità: {tematica}; argomento: {sottotematica}", level=1)
    
    for esercizio in verifica.esercizi:
        titolo = f"Esercizio numero {i}"
        testo = esercizio.testo
        doc.add_heading(titolo, level=2)

        doc.add_paragraph(testo)

        i += 1

    Cartella_controllo(file_path)

    doc.save(f"{file_path}Verifica_{classificazione}.docx")


def Scrivi_verifica_d(verifica: Verifica, classificazione: str):
    """
    Scrive la versione utile al docente della verifica in un file .docx.

    Args:
        verifica (Verifica): Oggetto Verifica contenente gli esercizi della verifica.
        classificazione (str): La classificazione della verifica.

    Returns:
        None
    """
    tematica = verifica.esercizi[0].tematica
    sottotematica = verifica.esercizi[0].argomento.sottotematica
    materia = verifica.esercizi[0].materia
    conteggio_infamia = 0
    i = 1
    file_path = f"./Smart_Test/Verifiche_{materia}/"
   

    doc = Document()

    doc.add_heading(f"Verifica di {materia}; unità: {tematica}; argomento: {sottotematica}", level=1)
    
    for esercizio in verifica.esercizi:
        titolo = f"Esercizio numero {i}"
        testo = esercizio.testo
        oda = esercizio.argomento.oda
        tipologia = esercizio.difficolta.tipologia
        livello = esercizio.difficolta.livello
        conteggio_infamia += int(esercizio.difficolta.infamia)
        risposta = esercizio.risposta
        centralita = esercizio.argomento.centralita


        
        doc.add_heading(titolo, level=2)

        if esercizio.difficolta.dsa:
            p_dsa = doc.add_paragraph("Quesito particolarmente complesso per studenti con disturbi specifici dell'apprendimento", level=1)
            Colore_rosso(p_dsa)
        if esercizio.argomento.trasversalita:
            p_tras = doc.add_paragraph("Il quesito ha elementi di trasversalità con altre materie", level=1)
            Colore_rosso(p_tras)
        
        p_car = doc.add_paragraph(f"Tipologia esercizio: {tipologia}, Livello esercizio: {livello}\n Indicazioni sulla centralità del quesito rispetto all'argomento della verifica: {centralita}")
        Colore_rosso(p_car)

        p_oda = doc.add_paragraph(f"Obiettivi di apprendimento: \n{oda} \n")
        Colore_rosso(p_oda)

        doc.add_paragraph(testo)

        doc.add_heading("Risposta", level = 2).font.color.rgb = RGBColor(255, 0, 0)
        p_risp = doc.add_paragraph(risposta)
        Colore_rosso(p_risp)

        i += 1

    doc.add_heading(f"Infamia della verifica: {conteggio_infamia}")
    Cartella_controllo(file_path)

    doc.save(f"{file_path}Verifica_docente_{classificazione}.docx")